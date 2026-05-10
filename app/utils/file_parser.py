"""
File parser — extract plain text from uploaded files for RAG indexing.

Supported text formats: .md, .txt, .pdf, .docx
Other formats (.jpg, .png, .json, .zip, etc.) return None — blob is stored
but not indexed in Qdrant.
"""
from pathlib import Path
from loguru import logger


# Mapping: extension → file_type category for the FE
DOCUMENT_EXTS = {"md", "txt", "pdf", "docx"}
IMAGE_EXTS = {"jpg", "jpeg", "png", "gif", "webp", "svg"}
DATA_EXTS = {"json", "csv", "xlsx", "xls"}

# Extensions whose text content can be extracted and embedded in Qdrant
INDEXABLE_EXTS = {"md", "txt", "pdf", "docx"}

EXT_TO_MIME = {
    "md": "text/markdown",
    "txt": "text/plain",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "gif": "image/gif",
    "webp": "image/webp",
    "svg": "image/svg+xml",
    "json": "application/json",
    "csv": "text/csv",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
    "zip": "application/zip",
}


def get_extension(filename: str) -> str:
    """Return lowercase extension without leading dot. Empty string if none."""
    return Path(filename).suffix.lstrip(".").lower()


def categorize(extension: str) -> str:
    """Return file_type category: document | image | data | other."""
    if extension in DOCUMENT_EXTS:
        return "document"
    if extension in IMAGE_EXTS:
        return "image"
    if extension in DATA_EXTS:
        return "data"
    return "other"


def get_mime_type(extension: str) -> str:
    return EXT_TO_MIME.get(extension, "application/octet-stream")


def is_indexable(extension: str) -> bool:
    return extension in INDEXABLE_EXTS


def extract_text(storage_path: str, extension: str) -> str | None:
    """
    Extract plain text from a saved file based on its extension.
    Returns None if the format is not indexable or extraction fails.
    """
    if not is_indexable(extension):
        return None

    try:
        if extension in {"md", "txt"}:
            return _read_text_file(storage_path)
        if extension == "pdf":
            return _extract_pdf(storage_path)
        if extension == "docx":
            return _extract_docx(storage_path)
    except Exception as e:
        logger.error(f"FileParser: Failed to extract text from {storage_path}: {e}")
        return None

    return None


def _read_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # also include tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
